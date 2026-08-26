#!/usr/bin/env python3
from __future__ import annotations

import argparse
import shutil
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NS = {'w': W_NS}


def set_update_fields_on_open(docx_path: Path) -> None:
    tmp = Path(tempfile.mkdtemp(prefix='docx_toc_patch_'))
    try:
        with zipfile.ZipFile(docx_path, 'r') as zf:
            zf.extractall(tmp)
        settings = tmp / 'word' / 'settings.xml'
        parser = etree.XMLParser(remove_blank_text=False)
        tree = etree.parse(str(settings), parser)
        root = tree.getroot()
        node = root.find('w:updateFields', namespaces=NS)
        if node is None:
            node = etree.Element(f'{{{W_NS}}}updateFields')
            root.insert(0, node)
        node.set(f'{{{W_NS}}}val', 'true')
        tree.write(str(settings), xml_declaration=True, encoding='UTF-8', standalone='yes')
        with zipfile.ZipFile(docx_path, 'w', compression=zipfile.ZIP_DEFLATED) as zf:
            for path in tmp.rglob('*'):
                if path.is_file():
                    zf.write(path, path.relative_to(tmp).as_posix())
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs)[::-1]:
        paragraph._p.remove(run._r)


def insert_toc_field(paragraph, levels: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    fld_begin.set(qn('w:dirty'), 'true')

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = f' TOC \\o "{levels}" \\h \\z \\u '

    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')

    placeholder = OxmlElement('w:t')
    placeholder.text = 'Inhoudsopgave wordt bijgewerkt...'

    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')

    run._r.extend([fld_begin, instr, fld_sep, placeholder, fld_end])


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument('docx', type=Path)
    ap.add_argument('--placeholder', default='[[TOC]]')
    ap.add_argument('--levels', default='1-3')
    args = ap.parse_args()

    doc = Document(str(args.docx))
    target = next((p for p in doc.paragraphs if (p.text or '').strip() == args.placeholder), None)
    if target is None:
        raise RuntimeError(f'TOC placeholder not found: {args.placeholder}')
    clear_paragraph(target)
    insert_toc_field(target, args.levels)
    doc.save(str(args.docx))
    set_update_fields_on_open(args.docx)
    print(f'PASS: genuine Word TOC field inserted into {args.docx}')


if __name__ == '__main__':
    main()
